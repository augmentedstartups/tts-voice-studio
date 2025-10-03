import io
from pathlib import Path
from typing import Dict, Any
import PyPDF2
import docx

async def extract_text_from_file(file_content: bytes, filename: str) -> Dict[str, Any]:
    """
    Extract text from uploaded file
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        
    Returns:
        Dict with extracted text and metadata
    """
    file_extension = Path(filename).suffix.lower()
    
    if file_extension == ".txt":
        text = file_content.decode("utf-8")
        word_count = len(text.split())
        
        return {
            "text": text,
            "metadata": {
                "filename": filename,
                "format": "txt",
                "wordCount": word_count,
                "characterCount": len(text)
            }
        }
    
    elif file_extension == ".pdf":
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            
            text = "\n\n".join(text_parts)
            word_count = len(text.split())
            
            return {
                "text": text,
                "metadata": {
                    "filename": filename,
                    "format": "pdf",
                    "pages": len(pdf_reader.pages),
                    "wordCount": word_count,
                    "characterCount": len(text)
                }
            }
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    elif file_extension in [".docx", ".doc"]:
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            text = "\n\n".join(text_parts)
            word_count = len(text.split())
            
            return {
                "text": text,
                "metadata": {
                    "filename": filename,
                    "format": "docx",
                    "paragraphs": len(doc.paragraphs),
                    "wordCount": word_count,
                    "characterCount": len(text)
                }
            }
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def chunk_text(text: str, max_chars: int = 4000) -> list[str]:
    """
    Split text into chunks for processing
    
    Args:
        text: Text to split
        max_chars: Maximum characters per chunk
        
    Returns:
        List of text chunks
    """
    if len(text) <= max_chars:
        return [text]
    
    chunks = []
    paragraphs = text.split("\n\n")
    current_chunk = ""
    
    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 <= max_chars:
            current_chunk += para + "\n\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = para + "\n\n"
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

